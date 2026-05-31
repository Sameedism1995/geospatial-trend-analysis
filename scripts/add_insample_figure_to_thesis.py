#!/usr/bin/env python3
"""Insert in-sample ML feature-importance figure into the revised thesis docx."""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
THESIS = ROOT / "outputs/thesis_strengthening/Thesis_Final_Draft_revised.docx"
FIGURE = ROOT / "outputs/final_figures/fig_ml_insample_feature_importance.png"

ANCHOR = (
    "The prediction of absolute future NDTI values was shown to be much more difficult "
    "than the modelling of the behaviour of temporal NDTI change."
)
ALREADY_MARKER = "In-sample feature importance · training weeks only"

INTRO = (
    "Although temporally separated test R² values were negative (Table above), in-sample "
    "feature-importance diagnostics on the training partition remain a legitimate "
    "explainability contribution: they show which spatial, maritime, and spectral inputs "
    "the models weighted most heavily when fit to the first 75% of calendar weeks. "
    "Figure 5.11 summarises Ridge standardized coefficients and HistGradientBoosting "
    "permutation importance computed on training rows only—not out-of-sample validation."
)

CAPTION = (
    "Figure 5.11. In-sample feature importance for ΔNDTI prediction (training weeks only). "
    "Top panels: Ridge absolute standardized coefficients (left) and HistGradientBoosting "
    "permutation importance ΔRMSE on a training subsample (right). Bottom panels isolate "
    "spatial (`grid_centroid_lat/lon`, `grid_res_deg`) and maritime (`vessel_density_t` lags) "
    "predictors; bar colours encode feature family. Weights reflect in-sample fit "
    "diagnostics and must not be read as forecast skill under temporal holdout (§5.9–5.10)."
)


def insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def insert_figure_block() -> bool:
    if not THESIS.is_file():
        print(f"Missing thesis: {THESIS}", file=sys.stderr)
        return False
    if not FIGURE.is_file():
        print(f"Missing figure: {FIGURE}", file=sys.stderr)
        return False

    doc = Document(str(THESIS))
    anchor_idx = None
    for i, para in enumerate(doc.paragraphs):
        if ANCHOR.split(".")[0] in para.text and "NDTI change" in para.text:
            anchor_idx = i
            break
    if anchor_idx is None:
        print("Anchor paragraph not found in thesis.", file=sys.stderr)
        return False

    # Idempotent: skip if already inserted
    for para in doc.paragraphs[anchor_idx : anchor_idx + 8]:
        if ALREADY_MARKER in para.text or "Figure 5.11" in para.text:
            print("Figure already present in thesis; skipping.")
            return True

    anchor = doc.paragraphs[anchor_idx]
    p_intro = insert_paragraph_after(anchor, INTRO)
    p_img = insert_paragraph_after(p_intro)
    p_img.alignment = 1  # center
    run = p_img.add_run()
    run.add_picture(str(FIGURE), width=Inches(6.4))
    insert_paragraph_after(p_img, CAPTION)

    doc.save(THESIS)
    print(f"Inserted Figure 5.11 into {THESIS}")
    return True


def main() -> int:
    return 0 if insert_figure_block() else 1


if __name__ == "__main__":
    raise SystemExit(main())
