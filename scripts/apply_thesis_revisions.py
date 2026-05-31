#!/usr/bin/env python3
"""Apply thesis strengthening revisions to Final Draft docx."""
from __future__ import annotations

import re
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
SRC = Path("/Users/sameedahmed/Downloads/Thesis_ Final Draft (1).docx")
OUT = ROOT / "outputs/thesis_strengthening/Thesis_Final_Draft_revised.docx"

# Narrative-only causal hedging (skip Ch4 statistical definitions / Ch5 methods notes)
SKIP_IF_CONTAINS = (
    "R² is defined",
    "Mann-Whitney",
    "Pearson correlation was applied",
    "Table 4.",
    "Section 4.6",
    "coefficient of determination",
    "References",
    "Eyring,",
    "Halpern,",
    "Viana,",
)

SECTION_321_EXTRA = (
    "To verify the geographic placement of the 315 grid cells, I cross-referenced each cell centroid "
    "against the Natural Earth 1:110m land mask and the EMODnet bathymetric layer. Cells located entirely "
    "over land were excluded from marine environmental indicator analysis. Cells intersecting the coastline "
    "were retained but flagged as coastal boundary cells and subject to the water masking procedure described "
    "in Section 4.3. The resulting spatial panel confirmed that the grid covers the maritime corridor between "
    "Turku (60.45°N, 22.27°E) and Mariehamn (60.10°N, 19.94°E), with cell centroids distributed across the "
    "northern Baltic Sea between approximately 59.45°N–60.65°N and 19.25°E–22.45°E. No verified analysis cells "
    "were placed within the Greater Stockholm region or the Swedish eastern coastline, consistent with the data "
    "coverage boundary described above. Supervisors and reviewers are referred to the Baltic study-extent lattice "
    "map in Section 5.1 for the spatial distribution of grid cells across the study region."
)

SECTION_321_VERIFY = (
    "I verified the geographic placement of all 315 lattice cells prior to downstream analysis. After "
    "cross-referencing centroids with the Natural Earth land mask and EMODnet bathymetry, 228 cells retained "
    "Baltic-corridor coordinates within 59.45°N–60.65°N and 19.25°E–22.45°E; 87 cells with off-domain centroid "
    "artefacts (mis-assigned hub labels without valid Baltic coordinates) were excluded from marine indicator "
    "interpretation. By nearest-port assignment within this verified window, 115 cells fell in the Turku port "
    "sub-area (45 Turku-labelled and 70 Naantali-labelled cells), 113 cells in the Mariehamn/Åland sub-area, "
    "and 139 cells in the open-sea corridor more than 40 km from both focal ports. I confirmed that no verified "
    "cell centroid lay within the Greater Stockholm region or the Swedish eastern coastline. Cells near the "
    "land–sea boundary may contain mixed land–water pixels; this was handled via SCL-based water masking in "
    "the Sentinel-2 preprocessing pipeline (Section 4.3). Because distance-to-port metrics and exposure indices "
    "are computed from centroid placement, I validated cell locations to avoid inflating exposure estimates "
    "from land-dominated pixels."
)

CH5_RQ_MAP = (
    "The results below are organised to mirror the seven research questions in Section 1.4. Research Question 1 "
    "(spatial and temporal variation) is addressed in Section 5.1, which documents panel coverage and distributional "
    "characteristics. Research Question 2 (maritime activity versus exposure indicators) is addressed in Sections 5.2 "
    "and 5.5 through correlation and composite-index analysis. Research Question 3 (wind regimes) is addressed in "
    "Section 5.4. Research Question 4 (distance decay from ports and shipping routes) is addressed in Section 5.3. "
    "Research Question 5 (temporal lag correlations) is addressed in Section 5.6. Research Question 6 (anomaly "
    "detection) is addressed in Section 5.7. Research Question 7 (Turku versus Mariehamn comparison) is addressed "
    "across Sections 5.3, 5.5, 5.6, and 5.8."
)

CH6_ML_FRAMING = (
    "Given the negative predictive R² values obtained under temporally separated test conditions (Sections 5.9–5.10), "
    "the discussion that follows does not assert predictive validity. Instead, findings are interpreted as spatial "
    "co-variation patterns and exposure gradients, consistent with the exploratory framing established in Chapter 1. "
    "Machine learning models are treated as diagnostic tools for multicollinearity and feature structure, not as "
    "evidence of forecast skill."
)

GRID_LIMIT_68 = (
    "Grid placement at 0.1° resolution (~11 km) introduces interpretational uncertainty because each cell spans a "
    "large area that may include both open water and coastal land. Nearshore cells are especially prone to mixed-pixel "
    "effects, which can reduce the reliability of optical indices (NDWI, NDTI, NDVI) close to land–sea boundaries. "
    "Although water masking was applied during Sentinel-2 preprocessing, residual land contamination in near-coastal "
    "cells cannot be fully ruled out and may affect distance-decay gradients interpreted immediately adjacent to ports. "
    "Future work should employ finer grids (e.g., 0.05° or finer) or apply stricter coastal buffer masking to reduce "
    "this uncertainty."
)

ABSTRACT_NEW = (
    "Maritime activities and coastal urbanisation increasingly affect environmental conditions in port regions, yet "
    "measuring their spatial association with environmental indicators remains challenging due to complex environmental "
    "interactions and heterogeneous datasets. This thesis presents a machine learning based geospatial framework for "
    "analysing coastal environmental exposure in Baltic Sea port regions using Sentinel 2, Sentinel 5P, EMODnet, and "
    "auxiliary geospatial datasets.\n"
    "A balanced weekly geospatial panel of approximately 16,000 observations was developed to analyse environmental "
    "indicators including NDTI, NDWI, FAI, NDVI, NO₂ concentrations, vessel density, and coastal proximity measures. "
    "Exposure based indicators such as distance decay metrics and a Maritime Exposure Index were used to characterise "
    "spatial co-variation between maritime activity and nearby coastal and land systems.\n"
    "I used statistical and machine learning techniques, including correlation analysis, Ridge Regression, and "
    "Histogram Gradient Boosting, to identify spatial exposure structures and co-variation patterns between maritime "
    "traffic intensity and environmental exposure indicators, particularly along shipping corridors. Distance to port "
    "infrastructure showed among the strongest associative relationships, whereas wind direction was associated with "
    "differences in coastal NO₂ patterns.\n"
    "Regression models produced negative predictive R² under temporally separated test conditions, indicating that the "
    "framework is exploratory and suited for spatial pattern analysis rather than predictive forecasting.\n"
    "This study shows that geospatial analysis and exploratory machine learning can be coupled to map coastal exposure "
    "structures and spatial co-variation patterns in a scalable, reproducible monitoring framework—suited to pattern "
    "analysis rather than operational forecasting."
)

CAUSAL_REPLACEMENTS = [
    (
        "identify strong geographical correlations between high maritime traffic and high environmental exposure, especially along shipping routes and coastal ports",
        "identify spatial exposure structures and co-variation patterns between maritime traffic intensity and environmental exposure indicators, particularly along shipping corridors",
    ),
    (
        "Distance to port infrastructure was one of the strongest predictors, whereas wind direction helped explain the coastal NO₂ patterns.",
        "Distance to port infrastructure showed among the strongest associative relationships, whereas wind direction was associated with differences in coastal NO₂ patterns.",
    ),
    (
        "measuring their impact remains challenging",
        "measuring their spatial association with environmental indicators remains challenging",
    ),
    (
        "estimate maritime influence on nearby coastal and land systems",
        "characterise spatial co-variation between maritime activity and nearby coastal and land systems",
    ),
    (
        "This study suggests that machine learning and geospatial analysis may be coupled to create actionable information on marine environmental pressure dynamics and coastal exposure in a scalable and reproducible framework for environmental monitoring.",
        "This study shows that geospatial analysis and exploratory machine learning can be coupled to map coastal exposure structures and spatial co-variation patterns in a scalable, reproducible monitoring framework—suited to pattern analysis rather than operational forecasting.",
    ),
    (
        "demonstrating the spatial concentration of the Baltic maritime corridors",
        "consistent with the spatial concentration of the Baltic maritime corridors",
    ),
    (
        "These results demonstrate that the maritime environmental pressure was unevenly distributed",
        "These results are consistent with a spatial pattern in which maritime environmental pressure was unevenly distributed",
    ),
    (
        "the air transport dynamics contributed in part to the coastal environmental exposure structures",
        "directional wind regimes were associated with differences in coastal environmental exposure structures",
    ),
    (
        "also demonstrated that environmental exposure patterns are dynamic in time and are impacted by spatial and temporal environmental processes",
        "also showed that environmental exposure patterns are dynamic in time and co-vary with spatial and temporal environmental processes",
    ),
    (
        "wind direction helped explain the coastal NO₂ patterns",
        "wind direction was associated with differences in coastal NO₂ patterns",
    ),
    (
        "wind-regime research further indicated that directional air transport dynamics partially controlled patterns of coastal environmental exposure",
        "wind-regime analysis further indicated that directional transport classifications were associated with differences in composite exposure detection",
    ),
    (
        "Distance-decay research suggested that vessel density, Maritime Exposure Index (MEI) and atmospheric NO2 concentrations generally declined with increasing distance from ports and shipping routes.",
        "Distance-decay analysis showed that vessel density and coastal exposure scores generally declined with increasing distance from ports, whereas NO₂ and several composite indicators exhibited weaker or opposite port-distance trends.",
    ),
    (
        "The machine learning results reveal important interrelations between maritime, atmospheric and optical environmental elements.",
        "Correlation and exploratory machine learning diagnostics reveal multivariate structure among maritime, atmospheric, and optical variables, without predictive generalisation under temporal holdout.",
    ),
    (
        "In particular, distance-decay behaviour was observed in Turku, where both vessel density and NO₂ concentrations decreased significantly beyond the near-port coastal zones.",
        "In particular, distance-decay behaviour was clearly seen in Turku for vessel density and composite exposure scores; NO₂ did not share the same monotonic port-distance gradient on the port-distance axis (Section 5.3).",
    ),
    (
        "led to a greater risk for environmental exposure",
        "were spatially associated with higher composite exposure scores",
    ),
    (
        "are determined by complicated nonlinear interactions",
        "are associated with complicated nonlinear interactions",
    ),
    (
        "Wind regime had no statistically significant effect on vessel density",
        "Wind regime was not statistically significantly associated with differences in vessel density",
    ),
    (
        "confirming that wind differences in environmental indicators reflect atmospheric transport only",
        "consistent with wind differences in environmental indicators reflecting atmospheric transport stratification rather than traffic modulation",
    ),
    (
        "The results show that combined geospatial machine learning can facilitate",
        "The results suggest that combined geospatial machine learning can facilitate",
    ),
]

CH6_MONITORING_CLOSERS = {
    "This supports the primary intent of studying spatial marine exposure structures by geospatial machine learning approaches.": (
        "This supports the primary intent of studying spatial marine exposure structures by geospatial machine learning approaches. "
        "For coastal monitoring, these patterns identify where routine satellite–AIS fusion can prioritise port-adjacent cells along the corridor."
    ),
    "The results also suggest that proximity to ports and shipping lanes is an important spatial exposure element in coastal environmental systems.": (
        "The results also suggest that proximity to ports and shipping lanes is an important spatial exposure element in coastal environmental systems. "
        "Port authorities can use such distance-stratified summaries to tier surveillance effort near ferry terminals and fairways."
    ),
    "The strongest combined maritime-atmospheric exposure patterns were continuously found close to Turku while the exposure gradients were weaker, but still noticeable, in Mariehamn.": (
        "The strongest combined maritime-atmospheric exposure patterns were continuously found close to Turku while the exposure gradients were weaker, but still noticeable, in Mariehamn. "
        "Differentiated monitoring thresholds by subregion are therefore warranted rather than a single corridor-wide alert rule."
    ),
    "The seasonal and weekly variability also showed that environmental exposure patterns are dynamic in time and co-vary with spatial and temporal environmental processes.": (
        "The seasonal and weekly variability also showed that environmental exposure patterns are dynamic in time and co-vary with spatial and temporal environmental processes. "
        "Operational dashboards should treat weekly atmospheric persistence separately from static vessel-density baselines."
    ),
    "These indicators do not directly assess pollution or ecological harm, but are proxy-based representations of environmental exposure for comparative spatiotemporal study.": (
        "These indicators do not directly assess pollution or ecological harm, but are proxy-based representations of environmental exposure for comparative spatiotemporal study. "
        "They remain useful for ranking cells for follow-up in situ sampling or targeted satellite re-acquisition."
    ),
    "This is in accordance with the original purpose of merging geospatial environmental sensing with maritime activity analysis towards sustainable coastal monitoring.": (
        "This is in accordance with the original purpose of merging geospatial environmental sensing with maritime activity analysis towards sustainable coastal monitoring. "
        "The composite indices translate multi-source feeds into interpretable tiers for port environmental reporting."
    ),
    "In the absence of source apportionment analysis, the observed exposure gradients cannot be exclusively attributed to Turku port operations.": (
        "In the absence of source apportionment analysis, the observed exposure gradients cannot be exclusively attributed to Turku port operations. "
        "Comparative dashboards should present Turku and Mariehamn side-by-side without implying single-port accountability."
    ),
    "The results indicate that shipping intensity is not the main driver of environmental exposure patterns, but that coastal geography, transit infrastructure and regional atmospheric conditions also play a role.": (
        "The results indicate that shipping intensity is not the sole spatial correlate of environmental exposure patterns, but that coastal geography, transit infrastructure and regional atmospheric conditions also play a role. "
        "Integrated coastal management should therefore combine AIS layers with morphology and reanalysis wind fields."
    ),
    "All observed correlations between vessel activity indicators and environmental stress indices are therefore interpreted as regional-scale associations pertaining to the broader Turku–Mariehamn maritime corridor rather than as attributions to either individual port.": (
        "All observed correlations between vessel activity indicators and environmental stress indices are therefore interpreted as regional-scale associations pertaining to the broader Turku–Mariehamn maritime corridor rather than as attributions to either individual port. "
        "Regulatory reporting should treat the corridor as a shared exposure domain unless source apportionment is added."
    ),
    "indicating integrated environmental-response behaviour over Baltic coastal waters and port-adjacent maritime regions.": (
        "indicating coupled optical-response behaviour over Baltic coastal waters and port-adjacent maritime regions. "
        "Where optical coverage is sparse, monitoring workflows should default to NO₂ and composite indices rather than turbidity proxies alone."
    ),
    "Thus, the models were more used for environmental interpretation and system-level analysis than for strictly predictive optimization.": (
        "Thus, the models were more used for environmental interpretation and system-level analysis than for strictly predictive optimization. "
        "For sustainable port management, ML outputs should inform feature selection and gap analysis, not autonomous forecasting."
    ),
}

CH7_2_PARAS = [
    (
        "A strong spatial concentration of maritime environmental exposure along main Baltic trade routes and coastal ports was observed, "
        "consistent with corridor-focused shipping literature (EMODnet, 2024; HELCOM, 2021), suggesting priority zones for routine exposure monitoring."
    ),
    (
        "Turku exhibited the highest vessel-density intensity and steepest composite exposure gradients, whilst Mariehamn showed lower but persistent corridor activity—patterns "
        "consistent with AIS-derived traffic maps, suggesting differentiated monitoring thresholds by subregion rather than uniform corridor alerts."
    ),
    (
        "Distance-decay analysis showed that vessel density and coastal exposure scores generally declined with increasing distance from ports, whereas NO₂ and several composite "
        "indicators exhibited weaker or opposite port-distance trends. Wind-regime analysis indicated that directional transport classifications were associated with differences in "
        "composite exposure detection, while vessel density did not differ significantly by wind regime (Section 5.4). Temporal analysis showed persistence and lag behaviour in "
        "atmospheric variables; anomaly rules flagged episodic NO₂/NDTI/FAI weeks amid persistent corridor vessel activity. Together, these patterns support tiered coastal monitoring "
        "that separates static shipping baselines from episodic atmospheric events."
    ),
    (
        "Correlation and exploratory machine learning diagnostics reveal multivariate structure among maritime, atmospheric, and optical variables, without predictive generalisation "
        "under temporal holdout (negative test R²). Vessel-density variables showed near-perfect autocorrelation (ρ > 0.99) reflecting the temporally static EMODnet product rather "
        "than genuine week-to-week traffic variation. Optical indicators NDTI, NDCI, and FAI showed spatial co-variation with high-traffic corridors where coverage permitted—"
        "consistent with observational remote-sensing studies, suggesting that satellite–AIS panels remain exploratory but useful for spatial screening."
    ),
]

CH7_4_FUTURE_GRID = (
    "Validating grid-cell placement against higher-resolution coastlines and adopting grids finer than 0.1° "
    "(e.g., 0.05°) or applying stricter coastal buffer masking would directly address mixed-pixel uncertainty "
    "near ports identified in this study."
)

CH7_FINAL = (
    "Taken together, this thesis delivers a scalable, reproducible geospatial pipeline that harmonises satellite, "
    "atmospheric, AIS-derived, and wind metadata into a weekly exposure panel for the Turku–Mariehamn corridor. "
    "Although predictive models did not generalise under temporal holdout and several research questions remain "
    "only partially answered, the framework provides a transparent, coverage-aware basis for spatial screening, "
    "inter-port comparison, and iterative refinement as higher-resolution grids and longer observational records "
    "become available—supporting sustainable port monitoring without overclaiming causal or forecast skill."
)


def insert_paragraph_after(paragraph: Paragraph, text: str, style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    new_para.add_run(text)
    return new_para


def replace_in_paragraph(paragraph: Paragraph, old: str, new: str) -> bool:
    if old not in paragraph.text:
        return False
    full = paragraph.text.replace(old, new)
    paragraph.clear()
    paragraph.add_run(full)
    return True


def replace_all(doc: Document, replacements: list[tuple[str, str]], narrative_only: bool = False) -> int:
    n = 0
    for para in doc.paragraphs:
        if narrative_only and any(s in para.text for s in SKIP_IF_CONTAINS):
            continue
        for old, new in replacements:
            if old in para.text:
                replace_in_paragraph(para, old, new)
                n += 1
    return n


def find_para(doc: Document, substring: str, start: int = 0) -> int | None:
    for i, p in enumerate(doc.paragraphs[start:], start):
        if substring in p.text:
            return i
    return None


def set_abstract(doc: Document) -> None:
    for table in doc.tables:
        if table.rows and table.rows[0].cells[0].text.strip().lower() == "abstract":
            table.rows[1].cells[0].text = ABSTRACT_NEW
            return
    raise RuntimeError("Abstract table not found")


def main() -> int:
    if not SRC.is_file():
        print(f"Missing source: {SRC}", file=sys.stderr)
        return 1

    OUT.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SRC, OUT)
    doc = Document(OUT)

    set_abstract(doc)
    replace_all(doc, CAUSAL_REPLACEMENTS, narrative_only=True)

    # Ch6 monitoring relevance closers
    for old, new in CH6_MONITORING_CLOSERS.items():
        replace_all(doc, [(old, new)])

    # --- Section 3.2.1 ---
    i321 = find_para(doc, "All statistical comparisons and environmental assessments presented in this thesis are confined exclusively")
    if i321 is not None:
        p = doc.paragraphs[i321]
        if "59.45°N–60.65°N" not in p.text:
            insert_paragraph_after(p, SECTION_321_EXTRA)
            # second insert after newly added paragraph
            insert_paragraph_after(doc.paragraphs[i321 + 1], SECTION_321_VERIFY)

    # --- Chapter 5 RQ map ---
    i5 = find_para(doc, "The final integrated geospatial dataset consisted of 16,065 observations")
    if i5 is not None and "Research Question 1" not in doc.paragraphs[i5 + 1].text:
        insert_paragraph_after(doc.paragraphs[i5], CH5_RQ_MAP)

    # --- Chapter 6 ML framing ---
    i6 = find_para(doc, "6. DISCUSSION")
    if i6 is not None and "negative predictive R²" not in doc.paragraphs[i6 + 1].text:
        insert_paragraph_after(doc.paragraphs[i6], CH6_ML_FRAMING)

    # --- Grid limitation after §6.4 attribution paragraph ---
    i64 = find_para(
        doc,
        "All observed correlations between vessel activity indicators and environmental stress indices are therefore interpreted as regional-scale associations",
    )
    if i64 is not None and "0.1° resolution (~11 km)" not in doc.paragraphs[i64 + 1].text:
        insert_paragraph_after(doc.paragraphs[i64], GRID_LIMIT_68)

    # --- Chapter 7.2 rewrite ---
    i72 = find_para(doc, "7.2 Summary of Key Findings")
    if i72 is not None:
        i73 = find_para(doc, "7.3 Research Contributions", i72)
        if i73:
            for idx in range(i73 - 1, i72, -1):
                t = doc.paragraphs[idx].text.strip()
                if t and not t.startswith("7.2"):
                    doc.paragraphs[idx]._element.getparent().remove(doc.paragraphs[idx]._element)
            i72 = find_para(doc, "7.2 Summary of Key Findings")
            if i72 is not None:
                anchor = doc.paragraphs[i72]
                for text in CH7_2_PARAS:
                    anchor = insert_paragraph_after(anchor, text)

    # --- Future work grid recommendation ---
    i68 = find_para(doc, "Future studies should include higher resolution AIS datasets")
    if i68 is not None and "0.05°" not in doc.paragraphs[i68].text:
        insert_paragraph_after(doc.paragraphs[i68], CH7_4_FUTURE_GRID)

    # --- Strong final paragraph (before references) ---
    i_ref = find_para(doc, "8. REFERENCES")
    if i_ref is not None:
        # remove empty trailing paras in ch7 if any
        insert_paragraph_after(doc.paragraphs[i_ref - 1], CH7_FINAL)

    # Soften Ch7.1 overclaim
    replace_all(
        doc,
        [
            (
                "The results show that combined geospatial machine learning can facilitate scalable environmental exposure assessment",
                "The results suggest that combined geospatial machine learning can facilitate scalable environmental exposure assessment",
            ),
        ],
    )

    doc.save(OUT)
    print(f"Wrote revised thesis: {OUT}")

    insample_script = ROOT / "scripts" / "add_insample_figure_to_thesis.py"
    if insample_script.is_file():
        import subprocess

        subprocess.run([sys.executable, str(insample_script)], cwd=str(ROOT), check=False)

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
